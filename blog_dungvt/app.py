from flask import Flask, render_template, request, Response, redirect
import db
from docx import Document
import pathlib
import os

app = Flask(__name__)


@app.route("/", methods = ["GET"]) 
def homepage():
    # Render file HTML và trả về cho trình duyệt
    # Truyền dữ liệu để render
    error = False
    posts = db.history()
    return render_template("index.html", posts= posts, error = error)


@app.route("/", methods = ["POST"])
def new_post():
    error = False
    title = (request.form["title"])
    content = (request.form["content"])
    if title and content: 
        post = (title,content)
        db.log(post)
        return redirect('/', 302)
    else:
        error = True
    posts = db.history()
    return render_template("index.html", posts= posts, error = error)


@app.route("/post/<id>")
def detail(id):
    error = False
    post = db.get_detail(id)
    return Response(render_template("editpost.html",post = post, error = error), status=404, mimetype="text/html")
    # return render_template("editpost.html",post = post, error = error)


@app.route("/post/<id>",methods = ["POST"])
def edit_post(id):
    error = False
    title = (request.form["title"])
    content = (request.form["content"])
    if title and content: 
        post = (title,content,id)
        db.update(post)
        posts = db.history()   
        return render_template("index.html", posts= posts, error = error)
    else:
        error = True
        post = db.get_detail(id)
        return render_template("editpost.html", post=post, error = error)


@app.route("/about")
def about():
    # Render file HTML và trả về cho trình duyệt
    # Truyền dữ liệu để render
    return render_template("about.html")


@app.route("/resignation-letter",methods=["GET"])
def resignation_letter():
    # Render file HTML và trả về cho trình duyệt
    # Truyền dữ liệu để render
    return render_template("resignation-letter.html")

@app.route("/resignation-letter", methods=["POST"])
def save_doc():
    error = False
    is_save = False
    current_path = os.path.dirname(os.path.abspath(__file__))
    fullname = request.form["fullname"]
    reason = request.form["reason"]
    if fullname and reason: 
        doc = Document()
        doc.add_heading("Resignation Letter")
        doc.add_paragraph(fullname)
        doc.add_paragraph(reason)
        file_name ='Resignation-Letter.docx'
        doc.save(os.path.join(current_path,'static',file_name))
        is_save = True
        return render_template("resignation-letter.html", error = error, is_save = is_save, file_name=file_name)
    else: 
        error = True
        return render_template("resignation-letter.html", error = error)
    

# Kiểm tra nếu là main script
if __name__ == "__main__":
    # Chạy Flask app
    # Bật debug để restart server mỗi khi có thay đổi trong mã
    app.run(debug=True)